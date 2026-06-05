"""Genera el documento Word (.docx) de la planificación."""

import io
import re
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


VERDE = RGBColor(0x2F, 0x6F, 0x4E)
GRIS = RGBColor(0x4B, 0x55, 0x63)


def planificacion_a_docx(materia, grado, contenido, objetivo="", config=None):
    """Devuelve los bytes de un .docx armado a partir del texto generado."""
    config = config or {}
    doc = Document()

    estilo = doc.styles["Normal"]
    estilo.font.name = "Calibri"
    estilo.font.size = Pt(11)

    # Membrete (si la docente cargó escuela / su nombre en Configuración)
    membrete = "  ·  ".join(
        v for v in (config.get("escuela"), config.get("docente")) if v
    )
    if membrete:
        m = doc.add_paragraph()
        m.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = m.add_run(membrete)
        run.font.size = Pt(10)
        run.font.color.rgb = GRIS

    # Encabezado
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run("Planificación de Clase")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = VERDE

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"{materia}  ·  {grado}")
    run.font.size = Pt(12)
    run.font.color.rgb = GRIS

    fecha = doc.add_paragraph()
    fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fecha.add_run(datetime.now().strftime("%d/%m/%Y"))
    run.font.size = Pt(9)
    run.font.color.rgb = GRIS

    doc.add_paragraph()  # espacio

    _volcar_markdown(doc, contenido)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _volcar_markdown(doc, texto):
    """Convierte el texto (con encabezados/listas markdown) en párrafos Word."""
    for linea in texto.splitlines():
        linea = linea.rstrip()
        if not linea.strip():
            doc.add_paragraph()
            continue

        # Encabezados markdown (#, ##, ###) o líneas tipo "1. Título"
        encabezado = re.match(r"^(#{1,6})\s+(.*)$", linea)
        numerado_titulo = re.match(r"^(\d+)\.\s+([A-ZÁÉÍÓÚÑ].*)$", linea.strip())

        if encabezado:
            nivel = min(len(encabezado.group(1)), 3)
            _agregar_titulo(doc, _limpiar(encabezado.group(2)), nivel)
        elif numerado_titulo and len(numerado_titulo.group(2)) < 60:
            _agregar_titulo(doc, _limpiar(linea.strip()), 2)
        elif re.match(r"^\s*[-*•]\s+", linea):
            cuerpo = re.sub(r"^\s*[-*•]\s+", "", linea)
            doc.add_paragraph(_limpiar(cuerpo), style="List Bullet")
        else:
            p = doc.add_paragraph()
            _agregar_texto_con_negritas(p, linea)


def _agregar_titulo(doc, texto, nivel):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.bold = True
    run.font.color.rgb = VERDE
    run.font.size = Pt(15 - (nivel - 1) * 2)


def _agregar_texto_con_negritas(parrafo, texto):
    """Respeta los **negritas** de markdown dentro de un párrafo."""
    partes = re.split(r"(\*\*[^*]+\*\*)", texto)
    for parte in partes:
        if parte.startswith("**") and parte.endswith("**"):
            run = parrafo.add_run(parte[2:-2])
            run.bold = True
        else:
            parrafo.add_run(parte)


def _limpiar(texto):
    """Quita marcas markdown sobrantes para los títulos."""
    return texto.replace("**", "").replace("*", "").strip()
